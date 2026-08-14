"""Programmatic builders for ingestion test documents."""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument


def build_simple_pdf(page_texts: list[str]) -> bytes:
    """
    Build a minimal valid PDF with extractable text on each page.

    Uses raw PDF syntax to avoid test-only dependencies beyond pypdf.
    """
    if not page_texts:
        page_texts = [""]

    font_obj_num = 3 + len(page_texts) * 2
    fixed_objects: list[bytes] = []

    for i, text in enumerate(page_texts):
        safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET".encode("ascii")
        content_num = 3 + i * 2
        page_num = content_num + 1
        content_obj = (
            f"{content_num} 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream\nendobj\n"
        )
        page_obj = (
            f"{page_num} 0 obj\n"
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Contents {content_num} 0 R "
            f"/Resources << /Font << /F1 {font_obj_num} 0 R >> >> >>\n"
            f"endobj\n"
        ).encode("ascii")
        fixed_objects.extend([content_obj, page_obj])

    kids = " ".join(f"{3 + i * 2 + 1} 0 R" for i in range(len(page_texts)))
    pages_obj = (
        f"2 0 obj\n<< /Type /Pages /Kids [{kids}] /Count {len(page_texts)} >>\nendobj\n"
    ).encode("ascii")
    catalog_obj = b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    font_obj = (
        f"{font_obj_num} 0 obj\n"
        f"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n"
        f"endobj\n"
    ).encode("ascii")

    all_objs = [catalog_obj, pages_obj, *fixed_objects, font_obj]
    output = [b"%PDF-1.4\n"]
    offsets = [0]
    pos = len(output[0])
    for obj in all_objs:
        offsets.append(pos)
        output.append(obj)
        pos += len(obj)

    xref_lines = ["xref", f"0 {len(offsets)}", "0000000000 65535 f "]
    for off in offsets[1:]:
        xref_lines.append(f"{off:010d} 00000 n ")
    xref_section = "\n".join(xref_lines).encode("ascii") + b"\n"
    trailer = (
        f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\n"
        f"startxref\n{pos}\n%%EOF\n"
    ).encode("ascii")
    return b"".join(output) + xref_section + trailer


def build_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build an in-memory DOCX file."""
    doc = DocxDocument()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r_idx, row in enumerate(table_rows):
            for c_idx, value in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = value
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_txt(content: str, encoding: str = "utf-8") -> bytes:
    return content.encode(encoding)
