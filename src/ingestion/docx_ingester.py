"""DOCX document ingestion."""

from __future__ import annotations

import logging
from io import BytesIO
from typing import Optional
from zipfile import BadZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from src.ingestion.base import DocumentIngester, FileInput
from src.ingestion.errors import CorruptedFileError, ExtractionFailureError
from src.ingestion.io_helpers import normalize_extracted_text, read_bytes
from src.ingestion.validation import DocumentValidator
from src.models.document import Document, DocumentPage, DocumentSourceType
from src.utils.config import IngestionConfig, get_settings

logger = logging.getLogger(__name__)


class DocxIngester(DocumentIngester):
    """Extract paragraph and table text from DOCX files."""

    def __init__(
        self,
        config: Optional[IngestionConfig] = None,
        validator: Optional[DocumentValidator] = None,
    ) -> None:
        self._config = config or get_settings().ingestion
        self._validator = validator or DocumentValidator(self._config)

    @property
    def supported_types(self) -> list[DocumentSourceType]:
        return [DocumentSourceType.DOCX]

    def ingest(self, source: FileInput, filename: str | None = None) -> Document:
        if not filename:
            raise CorruptedFileError("DOCX ingestion requires a filename.")

        data = read_bytes(source)
        self._validator.validate_before_ingestion(data, filename)

        try:
            docx = DocxDocument(BytesIO(data))
        except (PackageNotFoundError, BadZipFile, ValueError, KeyError) as exc:
            logger.exception("Failed to read DOCX '%s'", filename)
            raise CorruptedFileError(
                f"Unable to read DOCX file '{filename}'. The file may be corrupted or invalid."
            ) from exc
        except Exception as exc:
            logger.exception("Unexpected DOCX read failure for '%s'", filename)
            raise ExtractionFailureError(
                f"Failed to extract content from DOCX file '{filename}'."
            ) from exc

        content_parts: list[str] = []
        warnings: list[str] = []

        for paragraph in docx.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)

        table_count = 0
        for table in docx.tables:
            table_count += 1
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    content_parts.append(" | ".join(cells))

        if not content_parts:
            warnings.append("DOCX contains no extractable paragraph or table text.")

        raw_text = normalize_extracted_text("\n".join(content_parts))
        pages = [DocumentPage(page_number=1, text=raw_text)] if raw_text else []

        metadata = {
            "paragraph_count": str(len(docx.paragraphs)),
            "table_count": str(table_count),
        }

        document = Document(
            source_type=DocumentSourceType.DOCX,
            filename=filename,
            raw_text=raw_text,
            pages=pages,
            metadata=metadata,
            extraction_warnings=warnings,
        )
        self._validator.validate_extracted_document(document)
        return document
